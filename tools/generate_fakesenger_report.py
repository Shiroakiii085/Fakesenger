from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "Bao_cao_du_an_Fakesenger.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(13)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)

    for style_name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        s = doc.styles[style_name]
        s.font.name = "Times New Roman"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        s.font.size = Pt(size)
        s.font.bold = True


def configure_paragraph(p, line_spacing=1.3, first_indent=True):
    fmt = p.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_after = Pt(6)
    if first_indent:
        fmt.first_line_indent = Cm(0.75)


def add_text(doc, text="", bold=False, italic=False, align=None, first_indent=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    configure_paragraph(p, first_indent=first_indent)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        configure_paragraph(p, first_indent=False)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        configure_paragraph(p, first_indent=False)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "D9EAD3")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()
    set_doc_defaults(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("TRƯỜNG ĐẠI HỌC: ................................................\n").bold = True
    p.add_run("KHOA/BỘ MÔN: ....................................................").bold = True
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO DỰ ÁN CÔNG NGHỆ PHẦN MỀM")
    r.bold = True
    r.font.size = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ĐỀ TÀI: XÂY DỰNG ỨNG DỤNG CHAT REALTIME FAKESENGER")
    r.bold = True
    r.font.size = Pt(18)
    for _ in range(3):
        doc.add_paragraph()
    add_text(doc, "Giảng viên hướng dẫn: ................................................", first_indent=False)
    add_text(doc, "Sinh viên thực hiện: ....................................................", first_indent=False)
    add_text(doc, "Mã số sinh viên: ............................................................", first_indent=False)
    add_text(doc, "Lớp: .............................................................................", first_indent=False)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Năm học 2025 - 2026")
    page_break(doc)

    # Preliminary pages
    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_text(
        doc,
        "Em xin chân thành cảm ơn giảng viên đã hướng dẫn, góp ý và tạo điều kiện để em hoàn thành dự án "
        "xây dựng ứng dụng chat realtime Fakesenger. Trong quá trình thực hiện, em đã có cơ hội vận dụng "
        "kiến thức về phân tích yêu cầu, thiết kế cơ sở dữ liệu, phát triển giao diện, xây dựng API, bảo mật "
        "và triển khai phần mềm trên môi trường thực tế."
    )
    add_text(
        doc,
        "Báo cáo này tổng hợp quá trình phân tích, thiết kế, hiện thực và đánh giá hệ thống. Do thời gian và "
        "kinh nghiệm còn hạn chế, sản phẩm khó tránh khỏi thiếu sót. Em rất mong nhận được ý kiến đóng góp "
        "để hệ thống được hoàn thiện hơn trong các phiên bản tiếp theo."
    )
    add_heading(doc, "TÓM TẮT DỰ ÁN", 1)
    add_text(
        doc,
        "Fakesenger là một ứng dụng web chat realtime được xây dựng theo định hướng đồ án sinh viên nhưng "
        "đảm bảo có thể sử dụng thực tế. Hệ thống hỗ trợ chat 1:1, nhóm N-N, kênh thông báo 1:N, gửi ảnh, "
        "ghi âm, gọi video cá nhân và gọi video nhóm, quản lý thành viên, duyệt yêu cầu thêm thành viên, "
        "chỉnh sửa và gỡ tin nhắn, quản lý hồ sơ người dùng, đồng thời triển khai miễn phí trên Vercel kết hợp "
        "Supabase."
    )
    page_break(doc)

    add_heading(doc, "MỤC LỤC", 1)
    toc_rows = [
        ("1", "Giới thiệu đề tài", "4"),
        ("2", "Phân tích yêu cầu", "5"),
        ("3", "Phân tích và thiết kế hệ thống", "6"),
        ("4", "Kiến trúc và công nghệ sử dụng", "7"),
        ("5", "Thiết kế cơ sở dữ liệu", "8"),
        ("6", "Thiết kế giao diện và chức năng", "9"),
        ("7", "Hiện thực hệ thống", "10"),
        ("8", "Kiểm thử phần mềm", "11"),
        ("9", "Triển khai và vận hành", "12"),
        ("10", "Đánh giá, hạn chế và hướng phát triển", "13"),
        ("11", "Kết luận", "14"),
        ("12", "Phụ lục", "15"),
    ]
    add_table(doc, ["STT", "Nội dung", "Trang"], toc_rows, [2, 12, 2])
    page_break(doc)

    # 1. Introduction
    add_heading(doc, "1. GIỚI THIỆU ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Lý do chọn đề tài", 2)
    add_text(
        doc,
        "Trong bối cảnh trao đổi trực tuyến trở thành nhu cầu thường xuyên, các hệ thống nhắn tin thời gian "
        "thực như Messenger, Telegram hoặc Zalo là ví dụ điển hình của phần mềm có tính ứng dụng cao. Việc "
        "xây dựng một ứng dụng chat giúp sinh viên tiếp cận nhiều nội dung quan trọng của công nghệ phần mềm: "
        "phân tích yêu cầu, quản lý người dùng, xử lý dữ liệu đồng thời, bảo mật, realtime, đa phương tiện và "
        "triển khai lên môi trường internet."
    )
    add_heading(doc, "1.2. Mục tiêu", 2)
    add_bullets(doc, [
        "Xây dựng một ứng dụng web chat có thể sử dụng được trên desktop và mobile.",
        "Hỗ trợ các mô hình trò chuyện phổ biến: chat riêng, nhóm và kênh thông báo.",
        "Áp dụng kiến trúc frontend/backend rõ ràng, có cơ sở dữ liệu và API.",
        "Triển khai miễn phí trên nền tảng cloud phù hợp với sinh viên.",
        "Thực hành các nguyên tắc công nghệ phần mềm từ phân tích đến kiểm thử và bảo trì."
    ])
    add_heading(doc, "1.3. Phạm vi", 2)
    add_text(
        doc,
        "Dự án tập trung vào các tính năng cốt lõi của một hệ thống nhắn tin hiện đại, bao gồm xác thực, "
        "tin nhắn văn bản và đa phương tiện, realtime, quản trị nhóm, video call và giao diện responsive. "
        "Các nội dung nâng cao như mã hóa đầu cuối, thông báo đẩy đa nền tảng, phân tích dữ liệu lớn và hệ "
        "thống chống spam tự động chưa nằm trong phạm vi phiên bản hiện tại."
    )
    page_break(doc)

    # 2. Requirements
    add_heading(doc, "2. PHÂN TÍCH YÊU CẦU", 1)
    add_heading(doc, "2.1. Tác nhân", 2)
    add_table(doc, ["Tác nhân", "Mô tả"], [
        ("Khách chưa đăng nhập", "Có thể đăng ký, đăng nhập."),
        ("Người dùng", "Trao đổi tin nhắn, tìm kiếm người dùng, tham gia phòng, gọi video, cập nhật hồ sơ."),
        ("Quản trị viên phòng", "Thêm/xóa thành viên, duyệt yêu cầu, xóa phòng, xóa toàn bộ tin nhắn."),
    ], [4, 12])
    add_heading(doc, "2.2. Yêu cầu chức năng", 2)
    add_table(doc, ["Mã", "Yêu cầu"], [
        ("FR01", "Đăng ký và đăng nhập bằng email/password."),
        ("FR02", "Tạo chat 1:1, nhóm N-N và kênh 1:N."),
        ("FR03", "Gửi/nhận tin nhắn realtime."),
        ("FR04", "Gửi ảnh, tin nhắn âm thanh và hiển thị lịch sử cuộc gọi."),
        ("FR05", "Gỡ tin nhắn với bản thân hoặc với mọi người, sửa tin nhắn văn bản."),
        ("FR06", "Tìm kiếm người dùng và tìm kiếm cuộc trò chuyện."),
        ("FR07", "Quản lý hồ sơ, ảnh đại diện, trạng thái cá nhân."),
        ("FR08", "Gọi video 1:1 và gọi video nhóm."),
        ("FR09", "Gửi yêu cầu thêm thành viên và cho phép admin duyệt."),
        ("FR10", "Xóa phòng và xóa toàn bộ tin nhắn sau khi xác nhận."),
    ], [3, 13])
    add_heading(doc, "2.3. Yêu cầu phi chức năng", 2)
    add_bullets(doc, [
        "Hiệu năng: tin nhắn hiển thị nhanh nhờ optimistic UI và realtime subscription.",
        "Bảo mật: sử dụng Row Level Security để giới hạn dữ liệu theo người dùng và quyền admin.",
        "Khả dụng: giao diện responsive, dùng được trên desktop và mobile.",
        "Khả năng triển khai: ưu tiên free tier của Vercel và Supabase.",
        "Khả năng bảo trì: tách rõ API route, client component, kiểu dữ liệu và schema SQL."
    ])
    page_break(doc)

    # 3. Analysis and design
    add_heading(doc, "3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "3.1. Mô hình nghiệp vụ", 2)
    add_text(
        doc,
        "Người dùng sau khi đăng nhập có thể tham gia nhiều phòng trò chuyện. Một phòng có thể thuộc ba loại: "
        "direct, group hoặc channel. Direct chỉ dành cho hai người dùng; group cho phép mọi thành viên trao đổi; "
        "channel chỉ admin gửi nội dung. Người dùng thường có thể đề xuất thêm thành viên mới vào nhóm, còn "
        "admin là người phê duyệt."
    )
    add_heading(doc, "3.2. Các ca sử dụng chính", 2)
    add_table(doc, ["Use case", "Tiền điều kiện", "Kết quả"], [
        ("UC01 Đăng nhập", "Người dùng có tài khoản", "Truy cập màn hình chat."),
        ("UC02 Gửi tin nhắn", "Là thành viên phòng", "Tin nhắn được lưu và broadcast realtime."),
        ("UC03 Tạo nhóm", "Đã đăng nhập", "Phòng mới được tạo cùng danh sách thành viên."),
        ("UC04 Gửi yêu cầu thêm người", "Là thành viên nhóm", "Yêu cầu pending xuất hiện với admin."),
        ("UC05 Gọi video", "Thành viên phòng có camera/micro", "Thiết lập kết nối WebRTC."),
        ("UC06 Xóa toàn bộ tin nhắn", "Là admin", "Dữ liệu tin nhắn bị xóa sau xác nhận."),
    ], [4, 5, 7])
    add_heading(doc, "3.3. Luồng xử lý tin nhắn", 2)
    add_numbered(doc, [
        "Người dùng nhập nội dung và bấm gửi.",
        "Client hiển thị ngay bản tin local để giảm cảm giác trễ.",
        "API kiểm tra token và lưu tin nhắn vào bảng messages.",
        "Supabase Realtime phát sự kiện INSERT cho các thành viên phòng.",
        "Client thay bản tin local bằng bản tin từ server và cập nhật thời gian phòng."
    ])
    add_heading(doc, "3.4. Luồng xử lý video call", 2)
    add_text(
        doc,
        "Video call sử dụng WebRTC cho luồng media và Supabase Broadcast cho signaling. Khi người gọi bắt đầu, "
        "hệ thống tạo bản ghi call trong messages, phát tín hiệu invite, sau đó các bên trao đổi offer, answer "
        "và ICE candidate. Nếu kết nối trực tiếp thất bại, hệ thống hỗ trợ cấu hình TURN để relay media."
    )
    page_break(doc)

    # 4. Architecture
    add_heading(doc, "4. KIẾN TRÚC VÀ CÔNG NGHỆ SỬ DỤNG", 1)
    add_heading(doc, "4.1. Kiến trúc tổng thể", 2)
    add_text(
        doc,
        "Hệ thống áp dụng kiến trúc web client-server. Giao diện người dùng được xây dựng bằng Next.js/React. "
        "Các route backend nằm trong app/api của Next.js, đóng vai trò trung gian giữa client và Supabase. "
        "Supabase cung cấp xác thực, PostgreSQL, Storage và Realtime. Riêng cuộc gọi sử dụng WebRTC kết hợp "
        "STUN/TURN để truyền media."
    )
    add_table(doc, ["Lớp", "Công nghệ", "Vai trò"], [
        ("Frontend", "Next.js 15, React 19, CSS thuần, lucide-react", "Hiển thị giao diện, xử lý tương tác."),
        ("Backend", "Next.js API Routes", "Xác thực request, thao tác dữ liệu nghiệp vụ."),
        ("Database", "PostgreSQL trên Supabase", "Lưu hồ sơ, phòng, thành viên, tin nhắn."),
        ("Realtime", "Supabase Realtime", "Đồng bộ tin nhắn, phòng, yêu cầu thành viên."),
        ("Media", "Supabase Storage, MediaRecorder, WebRTC", "Ảnh, audio, video call."),
        ("Deploy", "Vercel + Supabase free tier", "Triển khai internet chi phí thấp."),
    ], [4, 5, 7])
    add_heading(doc, "4.2. Lý do lựa chọn công nghệ", 2)
    add_bullets(doc, [
        "Next.js giúp gộp frontend và backend trong một repo, thuận lợi cho sinh viên triển khai.",
        "Supabase cung cấp Auth, PostgreSQL, Storage và Realtime trong cùng hệ sinh thái.",
        "WebRTC là chuẩn trình duyệt cho truyền media thời gian thực.",
        "Vercel phù hợp với quy trình CI/CD đơn giản từ GitHub.",
    ])
    add_heading(doc, "4.3. Sơ đồ kiến trúc logic", 2)
    add_text(
        doc,
        "Người dùng -> Giao diện React -> API Routes Next.js -> Supabase Auth/PostgreSQL/Storage\n"
        "Người dùng A <-> Supabase Broadcast (signaling) <-> Người dùng B\n"
        "Người dùng A <-> WebRTC + STUN/TURN <-> Người dùng B",
        first_indent=False
    )
    page_break(doc)

    # 5. Database
    add_heading(doc, "5. THIẾT KẾ CƠ SỞ DỮ LIỆU", 1)
    add_heading(doc, "5.1. Các bảng chính", 2)
    add_table(doc, ["Bảng", "Chức năng"], [
        ("profiles", "Thông tin người dùng: email, tên hiển thị, avatar, trạng thái."),
        ("rooms", "Thông tin phòng chat: loại phòng, tên, người tạo."),
        ("room_members", "Quan hệ nhiều-nhiều giữa người dùng và phòng."),
        ("messages", "Tin nhắn text, image, audio, call; trạng thái gỡ/sửa/cuộc gọi."),
        ("message_hides", "Danh sách tin nhắn bị ẩn riêng theo người dùng."),
        ("member_requests", "Yêu cầu thêm thành viên chờ admin duyệt."),
    ], [4, 12])
    add_heading(doc, "5.2. Quan hệ dữ liệu", 2)
    add_bullets(doc, [
        "profiles 1-N rooms thông qua created_by.",
        "profiles N-N rooms thông qua room_members.",
        "rooms 1-N messages.",
        "messages N-N profiles thông qua message_hides.",
        "member_requests liên kết requester, target và room.",
    ])
    add_heading(doc, "5.3. Bảo mật dữ liệu", 2)
    add_text(
        doc,
        "Dự án bật Row Level Security cho toàn bộ bảng nghiệp vụ. Chính sách RLS kiểm soát ai được đọc phòng, "
        "ai được gửi tin nhắn, ai được xóa thành viên hoặc duyệt yêu cầu. Với kênh thông báo, chỉ admin có thể "
        "gửi tin; với tin nhắn, người gửi chỉ được sửa hoặc gỡ tin của chính mình; với xóa phòng và xóa toàn bộ "
        "tin nhắn, hệ thống chỉ cho phép admin thực hiện."
    )
    add_heading(doc, "5.4. Chỉ mục và trigger", 2)
    add_text(
        doc,
        "Các chỉ mục được tạo trên room_members.user_id, room_members.room_id, messages(room_id, created_at), "
        "member_requests(room_id, status, created_at) và chỉ mục tìm kiếm profile. Trigger tự cập nhật "
        "updated_at của rooms khi có tin nhắn mới, giúp danh sách hội thoại sắp xếp theo hoạt động gần nhất."
    )
    page_break(doc)

    # 6. UI and functions
    add_heading(doc, "6. THIẾT KẾ GIAO DIỆN VÀ CHỨC NĂNG", 1)
    add_heading(doc, "6.1. Nguyên tắc giao diện", 2)
    add_bullets(doc, [
        "Giao diện tối giản, ưu tiên khả năng quét nhanh thông tin.",
        "Danh sách phòng chat và khung tin nhắn có thanh cuộn riêng.",
        "Mobile hiển thị một màn hình chính để tránh thao tác rườm rà.",
        "Các thao tác nguy hiểm như xóa dữ liệu đều có bước xác nhận.",
    ])
    add_heading(doc, "6.2. Danh sách chức năng hiện có", 2)
    add_table(doc, ["Nhóm chức năng", "Mô tả"], [
        ("Tài khoản", "Đăng ký, đăng nhập, đổi tên hiển thị, trạng thái, avatar."),
        ("Phòng chat", "Chat 1:1, nhóm, kênh; tìm kiếm hội thoại; quản lý thành viên."),
        ("Tin nhắn", "Text, ảnh, audio, sửa, gỡ với bản thân, gỡ với mọi người."),
        ("Realtime", "Tin nhắn mới, phòng mới, yêu cầu thành viên, call signal."),
        ("Cuộc gọi", "1:1, gọi nhóm, cuộc gọi nhỡ, bị từ chối, thời lượng."),
        ("Quản trị", "Duyệt yêu cầu, xóa thành viên, xóa phòng, xóa toàn bộ tin nhắn."),
    ], [4, 12])
    add_heading(doc, "6.3. Một số quyết định UX", 2)
    add_text(
        doc,
        "Người dùng có thể bấm trực tiếp vào bubble tin nhắn để mở menu thao tác; menu tự đổi hướng khi ở gần "
        "cuối danh sách để tránh bị composer che. Khi đổi phòng, khung tin nhắn có skeleton loading nhằm giảm "
        "cảm giác chờ. Lịch sử cuộc gọi được hiển thị như một dạng message để người dùng theo dõi xuyên suốt "
        "mạch hội thoại."
    )
    page_break(doc)

    # 7. Implementation
    add_heading(doc, "7. HIỆN THỰC HỆ THỐNG", 1)
    add_heading(doc, "7.1. Cấu trúc mã nguồn", 2)
    add_table(doc, ["Thư mục/Tệp", "Vai trò"], [
        ("app/", "Trang chính, CSS toàn cục, API route."),
        ("components/chat-shell.tsx", "Client component trung tâm của giao diện chat."),
        ("lib/", "Kiểu dữ liệu và helper Supabase."),
        ("supabase/", "Schema và các file migration SQL."),
        (".env.local", "Biến môi trường local."),
    ], [5, 11])
    add_heading(doc, "7.2. Một số điểm hiện thực nổi bật", 2)
    add_bullets(doc, [
        "Optimistic UI cho gửi tin nhắn nhằm giảm độ trễ cảm nhận.",
        "Dedupe message để tránh lặp dữ liệu khi vừa có response API vừa có realtime event.",
        "Ẩn tin nhắn theo từng người dùng bằng bảng message_hides.",
        "Media upload lên Supabase Storage, client lưu public URL vào messages.",
        "Video call dùng nhiều RTCPeerConnection khi gọi nhóm.",
        "TURN được cấu hình qua biến môi trường để cải thiện NAT traversal.",
    ])
    add_heading(doc, "7.3. Các API tiêu biểu", 2)
    add_table(doc, ["API", "Phương thức", "Chức năng"], [
        ("/api/me", "GET/POST/PATCH", "Hồ sơ người dùng."),
        ("/api/rooms", "GET/POST", "Danh sách và tạo phòng."),
        ("/api/direct", "POST", "Tạo hoặc lấy chat riêng."),
        ("/api/rooms/[roomId]/messages", "GET/POST/DELETE", "Đọc, gửi, xóa toàn bộ tin."),
        ("/api/messages/[messageId]", "PATCH/DELETE", "Sửa tin, cập nhật call, gỡ tin."),
        ("/api/rooms/[roomId]/member-requests", "GET/POST", "Yêu cầu thêm thành viên."),
    ], [6, 3, 7])
    page_break(doc)

    # 8. Testing
    add_heading(doc, "8. KIỂM THỬ PHẦN MỀM", 1)
    add_heading(doc, "8.1. Chiến lược kiểm thử", 2)
    add_text(
        doc,
        "Dự án áp dụng kiểm thử thủ công theo kịch bản chức năng kết hợp kiểm tra kiểu dữ liệu bằng TypeScript. "
        "Vì hệ thống có realtime và media, nhiều trường hợp cần được kiểm thử trên nhiều thiết bị và nhiều tài "
        "khoản khác nhau."
    )
    add_heading(doc, "8.2. Bảng ca kiểm thử đại diện", 2)
    add_table(doc, ["Mã", "Kịch bản", "Kết quả mong đợi"], [
        ("TC01", "Đăng nhập đúng tài khoản", "Vào màn hình chat thành công."),
        ("TC02", "Gửi tin nhắn văn bản", "Tin hiển thị ngay và đồng bộ bên nhận."),
        ("TC03", "Tạo nhóm và gửi yêu cầu thêm thành viên", "Admin nhận yêu cầu realtime."),
        ("TC04", "Gỡ tin với mọi người", "Tin chuyển sang trạng thái đã gỡ ở tất cả thành viên."),
        ("TC05", "Gửi ảnh và ghi âm", "File hiển thị/phát lại được."),
        ("TC06", "Gọi video 1:1 trên hai thiết bị", "Hai chiều hình/tiếng hoạt động."),
        ("TC07", "Gọi video nhóm", "Nhiều ô video hiển thị theo lưới."),
        ("TC08", "Xóa toàn bộ tin nhắn", "Có xác nhận trước khi dữ liệu bị xóa."),
    ], [2, 7, 7])
    add_heading(doc, "8.3. Kết quả kiểm thử", 2)
    add_text(
        doc,
        "Các chức năng cốt lõi đã được kiểm tra trong quá trình phát triển. Một số vấn đề thực tế từng phát sinh "
        "như trùng tin nhắn, delay khi đổi phòng, lỗi join do schema mới, lỗi iOS khi media permission bị chặn và "
        "vấn đề kết nối iPhone-iPhone do thiếu TURN đã được phân tích và cải thiện qua các phiên bản."
    )
    page_break(doc)

    # 9. Deployment
    add_heading(doc, "9. TRIỂN KHAI VÀ VẬN HÀNH", 1)
    add_heading(doc, "9.1. Môi trường triển khai", 2)
    add_table(doc, ["Thành phần", "Nền tảng"], [
        ("Ứng dụng web", "Vercel"),
        ("Cơ sở dữ liệu", "Supabase PostgreSQL"),
        ("Xác thực", "Supabase Auth"),
        ("Lưu trữ file", "Supabase Storage"),
        ("Realtime", "Supabase Realtime"),
        ("TURN", "Metered TURN hoặc dịch vụ tương đương"),
    ], [5, 11])
    add_heading(doc, "9.2. Các biến môi trường", 2)
    add_bullets(doc, [
        "NEXT_PUBLIC_SUPABASE_URL",
        "NEXT_PUBLIC_SUPABASE_ANON_KEY",
        "NEXT_PUBLIC_TURN_URLS",
        "NEXT_PUBLIC_TURN_USERNAME",
        "NEXT_PUBLIC_TURN_CREDENTIAL",
    ])
    add_heading(doc, "9.3. Quy trình triển khai", 2)
    add_numbered(doc, [
        "Push mã nguồn lên GitHub.",
        "Import repository vào Vercel.",
        "Khai báo biến môi trường.",
        "Chạy các file SQL migration trên Supabase.",
        "Redeploy sau mỗi lần thay đổi schema hoặc biến môi trường.",
    ])
    add_heading(doc, "9.4. Vận hành", 2)
    add_text(
        doc,
        "Việc vận hành hệ thống tập trung vào theo dõi quota free tier, giám sát lỗi realtime/media, kiểm tra "
        "quyền database và cập nhật migration đồng bộ giữa source code với môi trường production."
    )
    page_break(doc)

    # 10. Evaluation
    add_heading(doc, "10. ĐÁNH GIÁ, HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_heading(doc, "10.1. Kết quả đạt được", 2)
    add_bullets(doc, [
        "Hoàn thành hệ thống chat realtime có thể sử dụng được.",
        "Có phân quyền rõ ràng giữa admin và member.",
        "Hỗ trợ đa phương tiện và gọi video.",
        "Triển khai được trên môi trường internet với chi phí thấp.",
        "Giao diện đã được tối ưu dần cho desktop và mobile."
    ])
    add_heading(doc, "10.2. Hạn chế", 2)
    add_bullets(doc, [
        "Chưa có mã hóa đầu cuối.",
        "Lịch sử tin nhắn mới giới hạn 200 bản ghi gần nhất trong API hiện tại.",
        "Chưa có push notification.",
        "Gọi video phụ thuộc chất lượng mạng và TURN server.",
        "Chưa có hệ thống logging, monitoring và rate limiting chuyên sâu."
    ])
    add_heading(doc, "10.3. Hướng phát triển", 2)
    add_bullets(doc, [
        "Bổ sung phân trang/lazy loading tin nhắn.",
        "Thêm thông báo đẩy và trạng thái đã xem.",
        "Bổ sung emoji, reply, pin message, mention.",
        "Nghiên cứu mã hóa đầu cuối và quản lý khóa.",
        "Tích hợp monitoring, analytics và test tự động.",
    ])
    page_break(doc)

    # 11. Conclusion
    add_heading(doc, "11. KẾT LUẬN", 1)
    add_text(
        doc,
        "Dự án Fakesenger đã đáp ứng mục tiêu xây dựng một ứng dụng chat realtime có kiến trúc rõ ràng, giao diện "
        "thực dụng và nhiều chức năng gần với sản phẩm thực tế. Quá trình thực hiện giúp củng cố kiến thức về "
        "phân tích yêu cầu, thiết kế cơ sở dữ liệu, bảo mật bằng RLS, lập trình frontend, API backend, realtime, "
        "media browser API và triển khai cloud."
    )
    add_text(
        doc,
        "Từ góc nhìn công nghệ phần mềm, dự án thể hiện đầy đủ vòng đời phát triển: xác định bài toán, đặc tả "
        "yêu cầu, thiết kế, hiện thực, kiểm thử, sửa lỗi, triển khai và đánh giá. Đây là nền tảng phù hợp để "
        "tiếp tục mở rộng thành một hệ thống nhắn tin hoàn chỉnh hơn trong tương lai."
    )
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_numbered(doc, [
        "Tài liệu Next.js App Router.",
        "Tài liệu Supabase Auth, Database, Storage và Realtime.",
        "Tài liệu MDN về MediaRecorder và WebRTC.",
        "Tài liệu Vercel về triển khai Next.js.",
        "Tài liệu Metered TURN về cấu hình ICE server.",
    ])
    page_break(doc)

    # 12. Appendix
    add_heading(doc, "12. PHỤ LỤC", 1)
    add_heading(doc, "12.1. Danh sách migration đã sử dụng", 2)
    add_bullets(doc, [
        "supabase/schema.sql",
        "supabase/add-member-requests.sql",
        "supabase/add-media-and-message-removal.sql",
        "supabase/add-room-destructive-actions.sql",
        "supabase/add-call-history.sql",
    ])
    add_heading(doc, "12.2. Cấu trúc route API", 2)
    add_bullets(doc, [
        "/api/direct",
        "/api/me",
        "/api/messages/[messageId]",
        "/api/profiles/search",
        "/api/rooms",
        "/api/rooms/[roomId]",
        "/api/rooms/[roomId]/messages",
        "/api/rooms/[roomId]/members",
        "/api/rooms/[roomId]/member-requests",
    ])
    add_heading(doc, "12.3. Gợi ý khi thuyết trình", 2)
    add_numbered(doc, [
        "Mở đầu bằng bài toán và lý do chọn đề tài.",
        "Trình bày 3 mô hình chat: 1:1, nhóm, kênh.",
        "Giải thích kiến trúc Vercel + Supabase + WebRTC.",
        "Demo realtime, gửi ảnh/audio, duyệt thành viên và video call.",
        "Kết thúc bằng đánh giá, khó khăn đã gặp và hướng phát triển.",
    ])

    # Footer
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Báo cáo dự án Fakesenger - Trang ")
        add_page_number(footer)

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
